from docx import Document

# Crear un nuevo documento de Word
doc = Document()
doc.add_heading('Resultados de Manhattan', 0)

# Añadir tabla con los datos proporcionados
data = [
    ['Mapa', 'Suma #exp', 'Suma #gen', 'Suma cost', 'Runtime total (s)'],
    ['Mapa 1', '41791', '27808', '2274', '1.48'],
    ['Mapa 2', '13753', '12202', '3361', '0.73'],
    ['Mapa 3', '2215', '2652', '706', '0.83'],
    ['Mapa 4', '87', '214', '79', '0.38'],
    ['Mapa 5', '687057', '108492', '4748', '19.17']
]

# Crear la tabla y añadir los datos
table = doc.add_table(rows=1, cols=len(data[0]))
table.style = 'Table Grid'

# Añadir los encabezados de la tabla
hdr_cells = table.rows[0].cells
for i, heading in enumerate(data[0]):
    hdr_cells[i].text = heading

# Añadir los datos a la tabla
for row in data[1:]:
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = value

# Guardar el documento
file_path = "/mnt/data/tabla_manhattan.docx"
doc.save(file_path)

file_path